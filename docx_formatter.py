
import os
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

def formatar_placeholders_calibri(doc):
    """
    Formatar todos os placeholders para fonte Calibri 11pt
    """
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if "{{" in run.text and "}}" in run.text:
                run.font.name = "Calibri"
                run.font.size = Pt(11)
    
    # Formatar placeholders em tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if "{{" in run.text and "}}" in run.text:
                            run.font.name = "Calibri"
                            run.font.size = Pt(11)

def adicionar_espacamento_imagens(doc):
    """
    Adicionar quebra de linha após cada imagem
    """
    for paragraph in doc.paragraphs:
        # Verificar se o parágrafo contém uma imagem
        if paragraph._element.xpath('.//pic:pic', namespaces={'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'}):
            # Adicionar espaçamento após a imagem
            paragraph.space_after = Pt(12)
            
            # Adicionar quebra de linha
            new_paragraph = paragraph.insert_paragraph_after()
            new_paragraph.space_after = Pt(6)

def obter_dimensoes_imagem(paragraph):
    """
    Obter dimensões da imagem em um parágrafo
    """
    try:
        drawing = paragraph._element.xpath('.//a:graphic', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
        if drawing:
            extent = paragraph._element.xpath('.//wp:extent', namespaces={'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'})
            if extent:
                # Converter EMU para CM (1 CM = 360000 EMU)
                largura_emu = int(extent[0].get('cx'))
                largura_cm = largura_emu / 360000
                return largura_cm
    except:
        pass
    return None

def criar_tabela_invisivel(doc, num_colunas):
    """
    Criar tabela invisível para organizar imagens em colunas
    """
    table = doc.add_table(rows=1, cols=num_colunas)
    
    # Remover bordas da tabela
    tbl = table._tbl
    tblPr = tbl.tblPr
    
    # Adicionar bordas invisíveis
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)
    
    return table

def organizar_imagens_por_colunas(doc):
    """
    Organizar imagens em colunas baseado no tamanho
    """
    paragrafos_com_imagens = []
    
    # Identificar parágrafos com imagens
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph._element.xpath('.//pic:pic', namespaces={'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'}):
            largura = obter_dimensoes_imagem(paragraph)
            paragrafos_com_imagens.append({
                'index': i,
                'paragraph': paragraph,
                'largura': largura
            })
    
    # Agrupar imagens por tamanho
    imagens_verticais_2col = []  # até 7,50 cm
    imagens_pequenas_3col = []   # até 5,92 cm
    imagens_horizontais = []     # largura completa
    
    for img_info in paragrafos_com_imagens:
        largura = img_info['largura']
        if largura and largura <= 5.92:
            imagens_pequenas_3col.append(img_info)
        elif largura and largura <= 7.50:
            imagens_verticais_2col.append(img_info)
        else:
            imagens_horizontais.append(img_info)
    
    # Processar imagens em grupos
    processar_grupo_imagens(doc, imagens_pequenas_3col, 3)
    processar_grupo_imagens(doc, imagens_verticais_2col, 2)
    
    # Imagens horizontais mantêm formatação original
    for img_info in imagens_horizontais:
        img_info['paragraph'].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def processar_grupo_imagens(doc, grupo_imagens, num_colunas):
    """
    Processar grupo de imagens organizando em tabelas
    """
    if not grupo_imagens:
        return
    
    # Agrupar imagens em lotes do tamanho das colunas
    for i in range(0, len(grupo_imagens), num_colunas):
        lote = grupo_imagens[i:i + num_colunas]
        
        # Encontrar posição de inserção (após a primeira imagem do lote)
        primeiro_paragraph = lote[0]['paragraph']
        posicao_insercao = None
        
        for j, p in enumerate(doc.paragraphs):
            if p == primeiro_paragraph:
                posicao_insercao = j
                break
        
        if posicao_insercao is not None:
            # Criar tabela invisível
            table = criar_tabela_invisivel(doc, num_colunas)
            
            # Mover imagens para células da tabela
            for idx, img_info in enumerate(lote):
                if idx < num_colunas:
                    cell = table.rows[0].cells[idx]
                    
                    # Copiar conteúdo da imagem para a célula
                    cell_paragraph = cell.paragraphs[0]
                    cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    # Mover elementos da imagem
                    for element in img_info['paragraph']._element:
                        cell_paragraph._element.append(element)
                    
                    # Limpar parágrafo original
                    img_info['paragraph'].clear()
            
            # Adicionar espaçamento após tabela
            table_paragraph = doc.paragraphs[posicao_insercao + 1]
            table_paragraph.space_after = Pt(12)

def aplicar_melhorias_formatacao(input_path, output_path):
    """
    Aplicar todas as melhorias de formatação solicitadas
    """
    print(f"Carregando documento: {input_path}")
    doc = Document(input_path)
    
    print("Aplicando formatação Calibri 11pt aos placeholders...")
    formatar_placeholders_calibri(doc)
    
    print("Adicionando espaçamento após imagens...")
    adicionar_espacamento_imagens(doc)
    
    print("Organizando imagens por colunas...")
    organizar_imagens_por_colunas(doc)
    
    print(f"Salvando documento formatado: {output_path}")
    doc.save(output_path)
    
    print("Formatação concluída com sucesso!")
    return True

def processar_arquivo_anexado(arquivo_path):
    """
    Processar arquivo .docx anexado aplicando melhorias
    """
    if not os.path.exists(arquivo_path):
        print(f"Erro: Arquivo não encontrado - {arquivo_path}")
        return False
    
    # Gerar nome do arquivo de saída
    base_name = os.path.splitext(os.path.basename(arquivo_path))[0]
    output_name = f"{base_name}_FORMATADO.docx"
    output_path = os.path.join("output", output_name)
    
    try:
        return aplicar_melhorias_formatacao(arquivo_path, output_path)
    except Exception as e:
        print(f"Erro ao processar arquivo: {str(e)}")
        return False
