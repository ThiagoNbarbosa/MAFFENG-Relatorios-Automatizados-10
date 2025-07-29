import os
import zipfile
import tempfile
import shutil
from datetime import datetime
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from PIL import Image, UnidentifiedImageError

# Folder processing order as specified
ORDEM_PASTAS = [
    "- Área externa", 
    "- Área interna", 
    "- Segundo piso",
    "- Detalhes",
    "- Vista ampla"
]

# Folders that should use normal text with bold instead of headings
PASTAS_TEXTO_NORMAL = ["- Detalhes", "- Vista ampla"]

def processar_zip(zip_path, dados_formulario):
    """
    Extract ZIP file and organize folder structure
    Returns structured content list for Word document insertion
    """
    print(f"Processing ZIP file: {zip_path}")
    
    # Create temporary directory for extraction
    with tempfile.TemporaryDirectory() as temp_dir:
        # Extract ZIP file
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)
        
        # Find the root folder (should be the only folder in temp_dir)
        extracted_items = os.listdir(temp_dir)
        if len(extracted_items) == 1 and os.path.isdir(os.path.join(temp_dir, extracted_items[0])):
            pasta_raiz = os.path.join(temp_dir, extracted_items[0])
        else:
            pasta_raiz = temp_dir
        
        # Process folder structure
        conteudo = []
        
        # Walk through directory structure
        for root, dirs, files in os.walk(pasta_raiz):
            # Sort directories according to specified order
            if root == pasta_raiz:
                dirs.sort(key=lambda x: (ORDEM_PASTAS.index(x) if x in ORDEM_PASTAS else len(ORDEM_PASTAS), x))
            
            # Calculate relative path and hierarchy level
            rel_path = os.path.relpath(root, pasta_raiz)
            if rel_path == '.':
                continue  # Skip root directory
            
            path_parts = rel_path.split(os.sep)
            nivel = len(path_parts) - 1
            
            # Add folder title based on hierarchy level
            folder_name = path_parts[-1]
            if nivel == 0:
                conteudo.append(folder_name)
            elif nivel == 1:
                conteudo.append(f"»{folder_name}")
            elif nivel == 2:
                conteudo.append(f"»»{folder_name}")
            else:
                conteudo.append(f"»»»{folder_name}")
            
            # Process images in current folder
            arquivos_imagens = [
                os.path.join(root, file)
                for file in files
                if file.lower().endswith(('.png', '.jpg', '.jpeg'))
            ]
            
            # Sort images by creation time
            arquivos_imagens.sort(key=lambda x: os.path.getctime(x) if os.path.exists(x) else 0)
            
            # Add images to content
            for imagem_path in arquivos_imagens:
                # Copy image to temporary location for processing
                import uuid
                unique_id = str(uuid.uuid4())[:8]
                temp_image_path = os.path.join(tempfile.gettempdir(), f"temp_img_{unique_id}_{os.path.basename(imagem_path)}")
                try:
                    shutil.copy2(imagem_path, temp_image_path)
                    print(f"Copied image: {os.path.basename(imagem_path)} -> {temp_image_path}")
                    conteudo.append({"imagem": temp_image_path})
                except Exception as e:
                    print(f"Error copying image {imagem_path}: {e}")
            
            # Add page break after each folder section
            if arquivos_imagens:  # Only add page break if there were images
                conteudo.append({"quebra_pagina": True})
        
        return conteudo

def substituir_placeholders(doc, dados_formulario, placeholders):
    """
    Replace placeholders in Word document with form data and apply specific formatting
    """
    print("Replacing placeholders in document...")
    
    # Create data mapping for placeholders
    placeholder_data = {}
    for placeholder, form_field in placeholders.items():
        if form_field in dados_formulario:
            placeholder_data[placeholder] = dados_formulario[form_field]
        elif form_field == 'Ygor Augusto Fernandes':
            placeholder_data[placeholder] = 'Ygor Augusto Fernandes'
        elif form_field == 'MAFFENG - Engenharia e Manutenção Profissional':
            placeholder_data[placeholder] = 'MAFFENG - Engenharia e Manutenção Profissional'
    
    # Replace placeholders in paragraphs with specific formatting
    for paragraph in doc.paragraphs:
        for placeholder, value in placeholder_data.items():
            if placeholder in paragraph.text:
                # Replace text preserving existing formatting and structure
                text_with_replacement = paragraph.text.replace(placeholder, str(value))
                
                # Clear runs but preserve paragraph structure
                for run in paragraph.runs:
                    run.clear()
                
                # Add new run with proper formatting
                run = paragraph.add_run(text_with_replacement)
                aplicar_estilo_placeholder(run, placeholder)
                
                # Ensure no hyperlink formatting is applied
                run.font.underline = False
    
    # Replace placeholders in tables with specific formatting
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, value in placeholder_data.items():
                        if placeholder in paragraph.text:
                            # Replace text preserving existing formatting and structure
                            text_with_replacement = paragraph.text.replace(placeholder, str(value))
                            
                            # Clear runs but preserve paragraph structure
                            for run in paragraph.runs:
                                run.clear()
                            
                            # Add new run with proper formatting
                            run = paragraph.add_run(text_with_replacement)
                            aplicar_estilo_placeholder(run, placeholder)
                            
                            # Ensure no hyperlink formatting is applied
                            run.font.underline = False
    
    print(f"Replaced {len(placeholder_data)} placeholders")

def aplicar_estilo(run, tamanho, negrito=False, fonte="Arial"):
    """Apply font styling to text run"""
    run.font.name = fonte
    run.font.size = Pt(tamanho)
    run.bold = negrito

def aplicar_estilo_placeholder(run, placeholder):
    """Apply specific font styling based on placeholder type"""
    # Arial PT 11 placeholders
    arial_placeholders = [
        '{{endereco_dependencia}}',
        '{{responsavel_dependencia}}',
        '{{responsavel_tecnico}}'
    ]
    
    # Calibri PT 11 placeholders  
    calibri_placeholders = [
        '{{ordem_servico}}',
        '{{data_elaboracao}}',
        '{{data_atendimento}}',
        '{{tipo_atendimento}}',
        '{{prefixo_sb}}',
        '{{nome_ag}}',
        '{{uf}}'
    ]
    
    if placeholder in arial_placeholders:
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.font.color.rgb = None  # Remove any color formatting
    elif placeholder in calibri_placeholders:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = None  # Remove any color formatting
    else:
        # Default formatting
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = None  # Remove any color formatting

def inserir_conteudo_word(modelo_path, conteudo, placeholders, dados_formulario, output_path):
    """
    Insert content into Word template and generate final document
    Returns number of images inserted
    """
    print(f"Loading Word template: {modelo_path}")
    
    # Count total images in content for tracking
    total_images = sum(1 for item in conteudo if isinstance(item, dict) and 'imagem' in item)
    print(f"Total images to process: {total_images}")
    
    # Load Word document
    doc = Document(modelo_path)
    contador_imagens = 0
    images_processed = 0
    
    # Replace placeholders first
    substituir_placeholders(doc, dados_formulario, placeholders)
    
    # Find insertion point
    paragrafo_insercao_index = None
    for i, paragrafo in enumerate(doc.paragraphs):
        if "{{start_here}}" in paragrafo.text:
            paragrafo_insercao_index = i
            # Clear the start_here marker
            paragrafo.text = paragrafo.text.replace("{{start_here}}", "")
            break
    
    if paragrafo_insercao_index is None:
        print("Warning: '{{start_here}}' marker not found in template")
        # Use last paragraph as insertion point
        paragrafo_insercao_index = len(doc.paragraphs) - 1
    
    # Insert content in reverse order to maintain proper positioning
    conteudo_invertido = list(reversed(conteudo))
    
    for item in conteudo_invertido:
        if isinstance(item, str):
            # Process folder titles
            titulo_limpo = item.replace("»", "").strip()
            # Remove duplicate hyphens if they exist
            if titulo_limpo.startswith("- -"):
                titulo_limpo = titulo_limpo[2:].strip()
            titulo = titulo_limpo + ":"
            nivel = item.count("»")
            
            # Insert new paragraph
            p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
            run = p.add_run(titulo)
            
            # Apply styling based on folder type and hierarchy
            if any(pasta in titulo for pasta in PASTAS_TEXTO_NORMAL):
                aplicar_estilo(run, 11, negrito=True)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            elif nivel == 0:
                p.style = 'Heading 1'
            elif nivel == 1:
                p.style = 'Heading 2'
            elif nivel == 2:
                p.style = 'Heading 3'
            else:
                aplicar_estilo(run, 12, negrito=True)
        
        elif isinstance(item, dict):
            if 'imagem' in item:
                # Process image insertion
                imagem_path = item["imagem"]
                images_processed += 1
                print(f"Processing image {images_processed}/{total_images}: {imagem_path}")
                
                # Validate image file exists and has content
                if not os.path.exists(imagem_path):
                    print(f"Error: Image file does not exist: {imagem_path}")
                    continue
                    
                if os.path.getsize(imagem_path) == 0:
                    print(f"Error: Image file is empty: {imagem_path}")
                    continue
                
                try:
                    # Validate image format first
                    with Image.open(imagem_path) as img:
                        img.verify()  # Verify image integrity
                    
                    # Re-open for actual processing (verify() closes the file)
                    with Image.open(imagem_path) as img:
                        largura_original, altura_original = img.size
                        formato = img.format
                        
                        print(f"Image details: {os.path.basename(imagem_path)} - {largura_original}x{altura_original} pixels, format: {formato}")
                        
                        # Convert image to RGB if necessary (for JPEG compatibility)
                        if img.mode not in ('RGB', 'L'):
                            img = img.convert('RGB')
                            # Save converted image to temporary file
                            temp_converted_path = imagem_path.replace('.png', '_converted.jpg')
                            img.save(temp_converted_path, 'JPEG', quality=95)
                            imagem_path = temp_converted_path
                            print(f"Converted image to RGB: {temp_converted_path}")
                        
                        altura_desejada_cm = 10  # Fixed height as specified
                        
                        # Calculate proportional width using correct DPI conversion
                        # Standard DPI conversion: 96 pixels = 2.54 cm (1 inch)
                        pixels_per_cm = 96 / 2.54  # approximately 37.8 pixels per cm
                        altura_original_cm = altura_original / pixels_per_cm
                        largura_original_cm = largura_original / pixels_per_cm
                        
                        # Calculate proportional width maintaining aspect ratio
                        aspect_ratio = largura_original_cm / altura_original_cm
                        largura_proporcional_cm = altura_desejada_cm * aspect_ratio
                        
                        # Limit maximum width to prevent page overflow
                        max_width_cm = 15  # Maximum width for document
                        if largura_proporcional_cm > max_width_cm:
                            largura_proporcional_cm = max_width_cm
                            altura_desejada_cm = max_width_cm / aspect_ratio
                        
                        print(f"Inserting image: {os.path.basename(imagem_path)} - {largura_proporcional_cm:.2f}cm x {altura_desejada_cm:.2f}cm")
                        
                        # Insert image paragraph
                        p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
                        
                        # Insert image
                        run = p.add_run()
                        run.add_picture(
                            imagem_path,
                            width=Cm(largura_proporcional_cm),
                            height=Cm(altura_desejada_cm)
                        )
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        contador_imagens += 1
                        
                        # Add single paragraph break after image for better spacing
                        p_break = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
                        
                        print(f"Successfully inserted image #{contador_imagens}: {os.path.basename(imagem_path)}")
                    
                    # Clean up temporary image files
                    try:
                        os.remove(imagem_path)
                        if 'converted' in imagem_path:
                            original_path = imagem_path.replace('_converted.jpg', '.png')
                            if os.path.exists(original_path):
                                os.remove(original_path)
                    except Exception as cleanup_error:
                        print(f"Warning: Could not clean up temporary file {imagem_path}: {cleanup_error}")
                
                except UnidentifiedImageError as e:
                    print(f"Error: Unrecognized or corrupted image format: {imagem_path} - {e}")
                except Exception as e:
                    print(f"Error inserting image '{imagem_path}': {e}")
                    import traceback
                    print(f"Full error traceback: {traceback.format_exc()}")
            
            elif 'quebra_pagina' in item:
                # Insert page break
                p = doc.paragraphs[paragrafo_insercao_index].insert_paragraph_before('')
                p.add_run().add_break(WD_BREAK.PAGE)
    
    # Save final document
    print(f"Saving document to: {output_path}")
    doc.save(output_path)
    
    print(f"Document created successfully with {contador_imagens} images")
    return contador_imagens
